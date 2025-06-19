import docx
from nltk.tokenize import sent_tokenize, word_tokenize
from nltk.corpus import stopwords
from nltk.probability import FreqDist
from nltk.tokenize.treebank import TreebankWordDetokenizer
from docx.shared import Pt
from nltk import pos_tag
from nltk.tokenize import RegexpTokenizer
from nltk.cluster.util import cosine_distance
from operator import itemgetter

def sentence_similarity(sent1, sent2):
    """Compute cosine similarity between two sentences."""
    all_words = list(set(sent1 + sent2))
    vector1 = [0] * len(all_words)
    vector2 = [0] * len(all_words)

    for w in sent1:
        vector1[all_words.index(w)] += 1

    for w in sent2:
        vector2[all_words.index(w)] += 1

    return 1 - cosine_distance(vector1, vector2)

def build_similarity_matrix(sentences):
    """Build a matrix representing the similarity between sentences."""
    matrix = [[sentence_similarity(sentences[i], sentences[j]) for j in range(len(sentences))]
              for i in range(len(sentences))]
    return matrix

def text_rank(sentence_matrix, num_sentences=3):
    """Apply TextRank algorithm to rank sentences."""
    scores = [1.0] * len(sentence_matrix)

    for _ in range(10):  # Iterative refinement (adjust as needed)
        new_scores = [0.15 + 0.85 * sum(sentence_matrix[i][j] * scores[j] for j in range(len(scores)))
                      for i in range(len(scores))]
        scores = new_scores

    ranked_sentences = sorted(enumerate(scores), key=itemgetter(1), reverse=True)
    selected_sentences = [sentence[0] for sentence in ranked_sentences[:num_sentences]]
    return selected_sentences

def text_summarization_with_text_rank(input_file, output_file, num_sentences=3):
    # Read the input text from the .docx file
    doc = docx.Document(input_file)
    input_text = ""
    for paragraph in doc.paragraphs:
        input_text += paragraph.text + " "

    # Tokenize the text into sentences
    sentences = [word_tokenize(sentence) for sentence in sent_tokenize(input_text)]

    # Calculate sentence similarity matrix
    sentence_matrix = build_similarity_matrix(sentences)

    # Apply TextRank algorithm to rank sentences
    selected_sentences = text_rank(sentence_matrix, num_sentences=num_sentences)

    # Create a new .docx file for the output summary
    output_doc = docx.Document()
    for index in selected_sentences:
        output_doc.add_paragraph(TreebankWordDetokenizer().detokenize(sentences[index]))

    # Save the output .docx file
    output_doc.save(output_file)

# Example usage
input_file = "C:\\21R21A6644\\Python\\input text.docx"
output_file = "C:\\21R21A6644\\Python\\output summary with TextRank.docx"
text_summarization_with_text_rank(input_file, output_file, num_sentences=10)
